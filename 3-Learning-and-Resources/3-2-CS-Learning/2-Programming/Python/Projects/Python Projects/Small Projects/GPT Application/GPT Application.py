import openai
import requests
import json
import re
import os
import random
from docx import Document

# 设置 OpenAI API Key 和 Google API Key
OPENAI_API_KEY = "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"
GOOGLE_API_KEY = "AIzaSyDdtlF6mDgiWUhEE4i_yn8jvjDkbl-GE-M"
GOOGLE_CSE_ID = "92ecc769947d44a95"
SEARCH_ENGINE_URL = "https://www.googleapis.com/customsearch/v1"
SERPAPI_KEY = "c23a6f9445cdeec3d4ba15e305e98328bcb3fffc8eec68faf07eb8c65783bf5e"
MODEL_NAME = "gpt-3.5-turbo-1106"

client = openai.OpenAI(api_key=OPENAI_API_KEY)


# ----------------------------------------
# Markdown Format
# ----------------------------------------


def format_outline_to_markdown(outline):
    lines = outline.split("\n")
    formatted_lines = []
    for line in lines:
        line = re.sub(r'^\d+(\.\d+)*\s+', '', line).strip()  # 删除所有编号并去除首尾空格
        line = re.sub(r'^\*\*(.*?)\*\*$', r'\1', line)  # 去掉行首尾的加粗 ** **
        
        heading_level = line.count('.')  # 计算层级（假设最多3级）
        line = '#' * heading_level + ' ' + line  # 添加对应的#号
        
        formatted_lines.append(line)
    return "\n".join(formatted_lines)


# ----------------------------------------
# Input Regulization
# ----------------------------------------

def translate_to_english(text):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": "You are a translator that translates Chinese text into English."},
            {"role": "user", "content": text}
        ],
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()  # Use dot notation



# ----------------------------------------
# Reference
# ----------------------------------------

def fetch_scholar_papers(query, num_results=25):
    url = "https://serpapi.com/search.json"
    params = {
        "engine": "google_scholar",
        "q": query,
        "api_key": SERPAPI_KEY,
        "num": num_results
    }
    response = requests.get(url, params=params)
    return response.json().get("organic_results", [])

    

def format_to_apa(paper, index):
    authors_list = paper.get("publication_info", {}).get("authors", [])
    formatted_authors = []
    for author in authors_list:
        name_parts = author["name"].split()
        if len(name_parts) > 1:
            formatted_authors.append(f"{name_parts[-1]}-{''.join([p[0] for p in name_parts[:-1]])}")
        else:
            formatted_authors.append(name_parts[0])
    authors = ", ".join(formatted_authors)
    
    year_match = re.search(r'\b(\d{4})\b', paper.get("publication_info", {}).get("summary", ""))
    year = year_match.group(1) if year_match else "n.d."
    
    title = paper.get("title", "Unknown Title")
    title = re.sub(r": (\w)", lambda m: f": {m.group(1).upper()}", title)
    
    summary_parts = paper.get("publication_info", {}).get("summary", "").split(" - ")
    journal = summary_parts[1] if len(summary_parts) > 1 else "Unknown Journal"
    
    doi = paper.get("link", "")
    
    return f"{authors}. ({year}). {title}. {journal}. DOI: {doi}", authors, year, index

def generate_in_text_citations(references):
    citations = []
    for ref in references:
        authors = ref[1].split(", ")
        year = ref[2]
        index = ref[3]
        if len(authors) == 1:
            citation = f"({authors[0]}, {year})[{index}]"
        elif len(authors) == 2:
            citation = f"({authors[0]} et {authors[1]}, {year})[{index}]"
        else:
            citation = f"({authors[0]} et al, {year})[{index}]"
        citations.append(citation)
    return citations

def insert_in_text_citations(text, citations):
    sentences = re.split(r'(\.|!|\?)', text)
    num_citations = min(len(sentences) // 3, len(citations))
    random_indices = random.sample(range(len(sentences)), num_citations)
    for i in sorted(random_indices):
        if not re.match(r'^\s*#+\s', sentences[i]) and citations:  # 确保不是标题
            sentences[i] = sentences[i].strip() + " " + citations.pop(0)
    return "".join(sentences)



# ----------------------------------------
# Main Body
# ----------------------------------------

def ask_chatgpt(prompt):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "system", "content": "你是一个智能联网搜索助手。"},
                  {"role": "user", "content": prompt}],
        temperature=0.7,
    )
    return response.choices[0].message.content


def generate_paper():
    topic_CN = input("Enter the research topic in Chinese: ")
    topic_CN = f'"{topic_CN}"'  # Add double quotes to the input
    topic = translate_to_english(topic_CN)

    prompt = f"""
    Generate a detailed academic outline (at least 3 levels) for the paper titled: 
    '{topic}'.
    The outline should include:
    1. Abstract (300-400 words)
    2. Key Words (3-4)
    3. Introduction (400-500 words)
    4. Literature Review (800-1000 words, at least 3 subsections)
    5. Methodology (detailed structure including empirical analysis)
    6. Results (including at least 5 tables, 3 images, statistical analysis)
    7. Discussion (500-1000 words, at least 3 aspects)
    8. Conclusion (max 500 words)
    9. References (APA format, 25 citations since 2021)
    for each section should have 3 levels eg 3.1, 3.1.1
"""
    outline = ask_chatgpt(prompt)
    formatted_outline = format_outline_to_markdown(outline)

    sections = {
        "Abstract": 350,
        "Key Words": 50,
        "Introduction": 450,
        "Literature Review": 900,
        "Methodology": 1000,
        "Results": 1000,
        "Discussion": 800,
        "Conclusion": 500
    }
    
    paper_content = {}
    for section, words in sections.items():
        if section in ["Results", "Discussion"]:
            paper_content[section] = ask_chatgpt(
                f"Write a {words}-word academic section on {section} for the topic: {topic}\n\n"
                f"Based on the following outline:\n{outline}, you should not include the section title at the beginning because we will add it later, "
                f"and you also can't add the number for subtitles like 4.1, 4.1.1. For those subtitles, you need to add ## before the level-2 titles "
                f"(like the original 4.1 in the outline) and ### before the level-3 titles (like the 4.2.2 in the outline). "
                f"For example, 4.1 xxx should be written as ## xxx. "
                f"Additionally, you should add 4+ tables and 3+ figures to provide quantitative support surrounding the {topic} with reasonable simulated data in this section."
            )
        else:
            paper_content[section] = ask_chatgpt(
                f"Write a {words}-word academic section on {section} for the topic: {topic}\n\n"
                f"Based on the following outline:\n{outline}, you should not include the section title at the beginning because we will add it later, "
                f"and you also can't add the number for subtitles like 4.1, 4.1.1. For those subtitles, you need to add ## before the level-2 titles "
                f"(like the original 4.1 in the outline) and ### before the level-3 titles (like the 4.2.2 in the outline). "
                f"For example, 4.1 xxx should be written as ## xxx."
            )

    
    papers = fetch_scholar_papers(topic)
    references = [format_to_apa(paper, i+1) for i, paper in enumerate(papers[:25])]
    in_text_citations = generate_in_text_citations(references)

        # 指定保存路径
    save_path = "/Users/evelyndu/Dropbox/缓冲资料库/dx"
    
    # 确保目录存在
    os.makedirs(save_path, exist_ok=True)
    
    # Markdown 文件路径
    md_file_path = os.path.join(save_path, "academic_paper.md")


    # 生成 Markdown 论文
    with open(md_file_path, "w", encoding="utf-8") as file:
        file.write(f"{formatted_outline}\n\n")
        for section, content in paper_content.items():
            updated_content = insert_in_text_citations(content, in_text_citations.copy())
            heading_level = 1 if section in ["Abstract", "Key Words", "Introduction", "Literature Review", 
                                             "Methodology", "Results", "Discussion", "Conclusion"] else 2
            file.write(f"{'#' * heading_level} {section}\n\n{updated_content}\n\n")
        file.write("# References\n\n")
        for ref in references:
            file.write(f"[{ref[3]}] {ref[0]}\n\n")
    
    print(f"Academic paper generated and saved as '{md_file_path}'.")
    
    # Word 文件路径
    docx_file_path = os.path.join(save_path, "academic_paper.docx")
    
    # 读取 Markdown 并转换为 Word
    doc = Document()
    with open(md_file_path, "r", encoding="utf-8") as file:
        for line in file:
            if line.startswith("# "):
                doc.add_heading(line.strip("# ").strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line.strip("## ").strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line.strip("### ").strip(), level=3)
            else:
                doc.add_paragraph(line.strip())
    
    doc.save(docx_file_path)
    print(f"Academic paper converted and saved as '{docx_file_path}'.")


# ----------------------------------------
# Execution
# ----------------------------------------

if __name__ == "__main__":
    generate_paper()